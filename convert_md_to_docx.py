"""
Convert Markdown files in docs/ folder to Word documents (.docx)
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def markdown_to_docx(md_file_path, docx_file_path):
    """Convert a markdown file to a Word document."""
    print(f"Converting {md_file_path} to {docx_file_path}...")
    
    # Read markdown file
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create Word document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip empty lines (but add spacing)
        if not line.strip():
            if i > 0 and lines[i-1].strip():  # Only add if previous line wasn't empty
                doc.add_paragraph()
            i += 1
            continue
        
        # Headers
        if line.startswith('# '):
            # H1
            p = doc.add_heading(line[2:].strip(), level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('## '):
            # H2
            p = doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            # H3
            p = doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('#### '):
            # H4
            p = doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith('##### '):
            # H5
            p = doc.add_heading(line[6:].strip(), level=5)
        elif line.startswith('###### '):
            # H6
            p = doc.add_heading(line[7:].strip(), level=6)
        
        # Code blocks
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(code_text)
                p.style = 'No Spacing'
                for run in p.runs:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Bullet lists
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:].strip()
            # Remove markdown formatting
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
            text = re.sub(r'`(.*?)`', r'\1', text)  # Code
            p = doc.add_paragraph(text, style='List Bullet')
        
        # Numbered lists
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            p = doc.add_paragraph(text, style='List Number')
        
        # Horizontal rules
        elif line.strip() == '---' or line.strip() == '***':
            doc.add_paragraph('_' * 50)
        
        # Regular paragraphs
        else:
            # Process inline formatting
            text = line
            # Remove markdown links but keep text
            text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
            # Remove code backticks
            text = re.sub(r'`([^`]+)`', r'\1', text)
            # Remove bold/italic but keep text
            text = re.sub(r'\*\*([^\*]+)\*\*', r'\1', text)
            text = re.sub(r'\*([^\*]+)\*', r'\1', text)
            
            if text.strip():
                p = doc.add_paragraph(text)
        
        i += 1
    
    # Save document
    doc.save(docx_file_path)
    print(f"Created {docx_file_path}")

def main():
    """Convert all markdown files in docs/ to Word documents."""
    # Get absolute path
    base_dir = Path(__file__).parent
    docs_dir = base_dir / 'docs'
    output_dir = base_dir / 'documents'
    
    # Create documents directory if it doesn't exist
    output_dir.mkdir(exist_ok=True)
    
    # Find all markdown files
    md_files = list(docs_dir.glob('*.md'))
    
    if not md_files:
        print("No markdown files found in docs/ folder")
        return
    
    print(f"Found {len(md_files)} markdown file(s) to convert\n")
    
    for md_file in md_files:
        docx_file = output_dir / f"{md_file.stem}.docx"
        try:
            markdown_to_docx(md_file, docx_file)
        except Exception as e:
            print(f"Error converting {md_file}: {e}")
            import traceback
            traceback.print_exc()
    
    print(f"\nConversion complete! Documents saved in {output_dir}/")

if __name__ == '__main__':
    main()

